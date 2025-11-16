from __future__ import annotations

import io
import json
import os
import subprocess
from datetime import datetime
from typing import Any, Dict, List, Optional

from duckduckgo_search import DDGS
from docx import Document
from openpyxl import Workbook

from app.database.connection import get_connection
from app.services.analytics import increment_counter

OLLAMA_MODEL = "gpt-oss:20b-cloud"


def ddg_sources(query: str, max_results: int = 5) -> List[Dict[str, str]]:
    try:
        with DDGS() as ddgs:
            results = ddgs.text(query, max_results=max_results)
            cleaned = []
            for result in results:
                href = result.get("href", "")
                if not href or "pdf" in href.lower():
                    continue
                cleaned.append(
                    {
                        "title": result.get("title", "Vacancy"),
                        "href": href,
                        "body": result.get("body", "")[:280],
                    }
                )
            return cleaned
    except Exception:
        return []


def should_use_ollama() -> bool:
    return os.getenv("HR_AGENT_USE_OLLAMA", "").lower() in {"1", "true", "yes"}


def call_llm(resume_text: str, vacancy_text: str, sources: List[Dict[str, str]]) -> Optional[Dict[str, Any]]:
    if not should_use_ollama():
        return None
    context = "\n".join(f"- {src['title']}: {src['body']}" for src in sources[:5])
    prompt = (
        "You are a senior HR analyst. Compare candidate resume and job description, using external context.\n"
        "You MUST respond with strict JSON containing keys:\n"
        "{\n"
        '  "match_score": int (0-100),\n'
        '  "strengths": ["..."],\n'
        '  "weaknesses": ["..."],\n'
        '  "skills_match": ["skill: assessment"],\n'
        '  "experience_assessment": "text",\n'
        '  "education_assessment": "text",\n'
        '  "development_plan": ["step"...],\n'
        '  "recommendations": ["..."],\n'
        '  "summary": "text"\n'
        "}\n"
        "Resume:\n"
        f"{resume_text}\n\n"
        "Vacancy:\n"
        f"{vacancy_text}\n\n"
        "Context:\n"
        f"{context or 'none'}\n"
        "Return JSON only."
    )
    try:
        result = subprocess.run(
            ["ollama", "run", OLLAMA_MODEL],
            input=prompt,
            text=True,
            capture_output=True,
            timeout=90,
            check=False,
        )
        raw = (result.stdout or "").strip()
        parsed = json.loads(raw if raw.startswith("{") else raw[raw.find("{") : raw.rfind("}") + 1])
        parsed["match_score"] = max(0, min(100, int(parsed.get("match_score", 0))))
        parsed["engine"] = "llm"
        return parsed
    except Exception:
        return None


def heuristic_analysis(resume_text: str, vacancy_text: str) -> Dict[str, Any]:
    resume_tokens = {token.lower() for token in resume_text.split() if len(token) > 3}
    vacancy_tokens = {token.lower() for token in vacancy_text.split() if len(token) > 3}
    overlap = sorted(vacancy_tokens & resume_tokens)
    missing = sorted(vacancy_tokens - resume_tokens)
    score = int((len(overlap) / max(1, len(vacancy_tokens))) * 100)
    score = max(10, min(90, score))
    strengths = [f"Упоминается {token}" for token in overlap[:5]] or ["Резюме содержит часть ключевых слов."]
    weaknesses = [f"Нет подтверждения навыка {token}" for token in missing[:5]] or ["Добавьте конкретные достижения."]
    skills_match = [f"{token}: strong fit" for token in overlap[:5]]
    return {
        "match_score": score,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "skills_match": skills_match,
        "experience_assessment": "Опыт частично совпадает с требованиями.",
        "education_assessment": "Образование следует уточнить в резюме.",
        "development_plan": ["Расширить раздел 'Проекты'", "Добавить конкретные результаты"],
        "recommendations": ["Добавить метрики эффективности", "Обновить резюме под конкретную вакансию"],
        "summary": "Рекомендовано улучшить резюме под требования вакансии.",
        "engine": "heuristic",
    }


def store_analysis(user_id: int, resume_excerpt: str, vacancy: str, result: Dict[str, Any], sources: List[Dict[str, str]]) -> int:
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO analyses (user_id, resume_excerpt, vacancy_description, result_json, sources_json, engine, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (
            user_id,
            resume_excerpt[:1200],
            vacancy[:1200],
            json.dumps(result, ensure_ascii=False),
            json.dumps(sources, ensure_ascii=False),
            result["engine"],
            datetime.utcnow().isoformat(),
        ),
    )
    conn.commit()
    analysis_id = cursor.lastrowid
    conn.close()
    increment_counter("analyses")
    if result["engine"] == "llm":
        increment_counter("llm_analyses")
    return int(analysis_id)


def fetch_analysis(user_id: int, analysis_id: int) -> Optional[Dict[str, Any]]:
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "SELECT * FROM analyses WHERE id = ? AND user_id = ?",
        (analysis_id, user_id),
    )
    row = cursor.fetchone()
    conn.close()
    if not row:
        return None
    record = dict(row)
    record["result_json"] = json.loads(record["result_json"])
    record["sources_json"] = json.loads(record.get("sources_json") or "[]")
    return record


def list_analyses(user_id: int, limit: int = 10) -> List[Dict[str, Any]]:
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, match_score, engine, created_at
        FROM (
            SELECT id,
                   json_extract(result_json, '$.match_score') AS match_score,
                   engine,
                   created_at
            FROM analyses
            WHERE user_id = ?
            ORDER BY created_at DESC
            LIMIT ?
        )
        """,
        (user_id, limit),
    )
    rows = cursor.fetchall()
    conn.close()
    return [
        {
            "id": row["id"],
            "match_score": int(row["match_score"] or 0),
            "engine": row["engine"],
            "created_at": row["created_at"],
        }
        for row in rows
    ]


def analysis_to_docx(payload: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading("HR Agent — отчёт анализа", level=1)
    doc.add_paragraph(f"Дата: {payload.get('created_at')}")
    doc.add_paragraph(f"Совпадение: {payload.get('match_score')}% ({payload.get('engine')})")
    doc.add_heading("Сильные стороны", level=2)
    for item in payload.get("strengths", []):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Зоны роста", level=2)
    for item in payload.get("weaknesses", []):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Навыки", level=2)
    doc.add_paragraph(", ".join(payload.get("skills_match", [])))
    doc.add_heading("План развития", level=2)
    for item in payload.get("development_plan", []):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Рекомендации", level=2)
    for item in payload.get("recommendations", []):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(f"Сводка: {payload.get('summary','')}")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def analysis_to_xlsx(payload: Dict[str, Any]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Analysis"
    ws.append(["Поле", "Значение"])
    ws.append(["Дата", payload.get("created_at")])
    ws.append(["Совпадение", f"{payload.get('match_score')}% ({payload.get('engine')})"])
    ws.append(["Сильные стороны", "; ".join(payload.get("strengths", []))])
    ws.append(["Зоны роста", "; ".join(payload.get("weaknesses", []))])
    ws.append(["Навыки", "; ".join(payload.get("skills_match", []))])
    ws.append(["План развития", "; ".join(payload.get("development_plan", []))])
    ws.append(["Рекомендации", "; ".join(payload.get("recommendations", []))])
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.read()
